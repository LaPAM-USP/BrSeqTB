#!/usr/bin/env python3

# ============================================================
# Clinical TB Resistance Report Generator (Word)
#
# Usage:
#   python3 clinicalReport.py <biosample ID>
#   python3 clinicalReport.py <biosample ID> --table input/input_table.xlsx
#
# Inputs:
#   - input/input_table.csv  (default)
#   - input/input_table.xlsx (optional with --table)
#   - results/qc_summary.xlsx
#   - results/resistance/<biosample>.xlsx
#   - database/omsCatalog/dictionary.xlsx
#   - database/clinicalTemplate/report_template.docx
#
# Output:
#   - clinicalReport/<biosample>.docx
#
# Description:
#   Generates a clinical Word report for tuberculosis drug
#   resistance. This script preserves the original logic and
#   formatting, performing only placeholder replacement and
#   report assembly.
# ============================================================

import sys
import os
import argparse
import pandas as pd
from datetime import datetime
from docx import Document
from collections import defaultdict

# ============================================================
# PROJECT DIR (script always lives in bin/)
# ============================================================

PROJECT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

# ============================================================
# CONFIGURATION
# ============================================================

INPUT_TABLE   = os.path.join(PROJECT_DIR, "input", "input_table.csv")
QC_SUMMARY    = os.path.join(PROJECT_DIR, "results", "qc_summary.xlsx")
RESISTANCE_DIR = os.path.join(PROJECT_DIR, "results", "resistance")
DICTIONARY    = os.path.join(PROJECT_DIR, "database", "omsCatalog", "dictionary.xlsx")
MODEL         = os.path.join(PROJECT_DIR, "assets", "templates", "report_template.docx")
OUT_DIR       = os.path.join(PROJECT_DIR, "results", "clinicalReport")


# ============================================================
# FIND & REPLACE — PARAGRAPHS / TABLES / HEADER / FOOTER
# ============================================================

def replace_in_paragraph(paragraph, mapping):
    for key, value in mapping.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(key, value)

def replace_in_cell(cell, mapping):
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, mapping)

def replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            replace_in_cell(cell, mapping)

def replace_header_footer(doc, mapping):
    for section in doc.sections:
        header = section.header
        footer = section.footer

        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph, mapping)
        for table in header.tables:
            replace_in_table(table, mapping)

        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph, mapping)
        for table in footer.tables:
            replace_in_table(table, mapping)

# ============================================================
# APPLY BOLD TO PLACEHOLDERS
# ============================================================

def bold_placeholder(doc, placeholder):
    for p in doc.paragraphs:
        for run in p.runs:
            if placeholder in run.text:
                run.bold = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if placeholder in run.text:
                            run.bold = True

# ============================================================
# HELPERS
# ============================================================

def fmt_date(v):
    try:
        return pd.to_datetime(v).strftime("%d/%m/%Y")
    except:
        return ""

def superscript(n):
    SUP = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")
    return str(n).translate(SUP)

def format_lineage(raw):
    if not raw:
        return ""
    items = [x.strip() for x in raw.split(";") if x.strip()]
    conv = []
    for item in items:
        if item.startswith("lineage"):
            conv.append("L" + item.replace("lineage", "", 1))
        else:
            conv.append(item)
    return "; ".join(conv)

def get_lineage(biosample):
    try:
        df = pd.read_excel(QC_SUMMARY, sheet_name="lineage", dtype=str).fillna("")
        row = df[df["biosample"] == biosample]
        if not row.empty:
            return format_lineage(row.iloc[0]["LINEAGE"])
    except:
        pass
    return ""

def resolve_input_table_path(table_path):
    """
    Resolve input table path for direct execution and Nextflow execution.

    Accepts:
    - input/input_table.csv
    - input/input_table.xlsx
    - input_table.csv
    - input_table.xlsx
    - absolute paths
    """

    # 1. As provided
    if os.path.exists(table_path):
        return table_path

    # 2. Relative to project directory
    project_relative = os.path.join(PROJECT_DIR, table_path)
    if os.path.exists(project_relative):
        return project_relative

    # 3. If Nextflow passed only basename, look inside PROJECT_DIR/input/
    input_relative = os.path.join(PROJECT_DIR, "input", os.path.basename(table_path))
    if os.path.exists(input_relative):
        return input_relative

    raise FileNotFoundError(
        f"Input table not found: {table_path}. "
        f"Tried: {table_path}, {project_relative}, {input_relative}"
    )


def read_input_table(table_path):
    table_path = resolve_input_table_path(table_path)

    ext = os.path.splitext(table_path)[1].lower()

    if ext == ".csv":
        encodings = ["utf-8-sig", "utf-16", "cp1252", "latin1"]

        last_error = None
        for enc in encodings:
            try:
                return pd.read_csv(table_path, dtype=str, encoding=enc).fillna("")
            except UnicodeDecodeError as e:
                last_error = e
                continue

        raise RuntimeError(
            f"Could not decode input table as CSV: {table_path}. "
            f"Last error: {last_error}"
        )

    if ext == ".xlsx":
        return pd.read_excel(table_path, dtype=str).fillna("")

    raise ValueError("Input table must have extension .csv or .xlsx")

# ============================================================
# TRANSLATE DRUG NAMES IN COVERAGE FAILURES
# ============================================================

def traduzir_farmaco_cov(cov_str, trad):
    if not cov_str:
        return cov_str

    itens = []
    for item in cov_str.split(";"):
        item = item.strip()
        if "_" in item:
            base, drug = item.rsplit("_", 1)
            drug_pt = trad.get(str(drug).strip().lower(), drug)
            itens.append(f"{base}_{drug_pt}")
        else:
            itens.append(item)

    return ";".join(itens)

# ============================================================
# RESISTANCE PROCESSING
# ============================================================

def process_resistance(biosample):

    df = pd.read_excel(f"{RESISTANCE_DIR}/{biosample}.xlsx").fillna("")
    df["drug"] = df["drug"].astype(str).str.lower().str.strip()
    df["evidence"] = df["evidence"].astype(str).str.lower().str.strip()

    trad = {
        "amikacin":"Amicacina","bedaquiline":"Bedaquilina","capreomycin":"Capreomicina",
        "clofazimine":"Clofazimina","ethambutol":"Etambutol","ethionamide":"Etionamida",
        "isoniazid":"Isoniazida","levofloxacin":"Levofloxacino","linezolid":"Linezolida",
        "moxifloxacin":"Moxifloxacino","pyrazinamide":"Pirazinamida",
        "rifampicin":"Rifampicina","streptomycin":"Estreptomicina",
        "kanamycin":"Canamicina","delamanid":"Delamanida"
    }

    try:
        df_dict = pd.read_excel(DICTIONARY, dtype=str).fillna("")
        dict_map = dict(zip(df_dict["comment"].str.strip(), df_dict["Tradução"].str.strip()))
    except:
        dict_map = {}

    df["comment"] = df["comment"].apply(
        lambda c: dict_map.get(str(c).strip(), str(c).strip())
    )

    df_r = df[df["evidence"] == "r"].copy()
    df_pass = df_r[df_r["filter_status"] == "PASS"].copy()

    borderline_set = {
        "rpoB_p.Leu430Pro","rpoB_p.Leu452Pro","rpoB_p.His445Tyr","rpoB_p.His445Leu",
        "rpoB_p.Asp435Tyr","rpoB_p.His445Asn","rpoB_p.His445Arg","rpoB_p.His445Cys"
    }

    comment_map = {}
    idx_counter = 1

    variants_by_drug = defaultdict(list)
    hetero = []
    borderline = []

    for _, row in df_pass.iterrows():
        var = row["variant"]
        drug = row["drug"]
        het = str(row["heteroresistance"]).upper() == "HET"
        comment = row["comment"].strip()

        idx = None
        if comment:
            if comment not in comment_map:
                comment_map[comment] = idx_counter
                idx_counter += 1
            idx = comment_map[comment]

        suffix = ""
        if het:
            suffix += "ʰ"
            try:
                af_fmt = f"{float(row['af']):.2f}"
            except:
                af_fmt = str(row["af"])
            alt_reads = str(row.get("alt_reads", "")).strip()
            hetero.append(f"{var} (AF: {af_fmt}; READS: {alt_reads})")

        if var in borderline_set:
            suffix += "†"
            borderline.append(var)

        if idx:
            suffix += superscript(idx)

        variants_by_drug[drug].append(f"{var}{suffix}")

    comentarios = "\n".join(f"{idx}. {text}" for text, idx in comment_map.items())

    resistant_drugs = sorted(set(df_pass["drug"]))
    all_drugs = sorted(set(df["drug"]))
    sensitive_drugs = sorted(d for d in all_drugs if d not in resistant_drugs)

    trad_r = "\n\n".join([trad[d] for d in resistant_drugs])
    var_str = "\n\n".join(", ".join(variants_by_drug[d]) for d in resistant_drugs)
    trad_s = "\n".join([trad[d] for d in sensitive_drugs])

    df_fail = df_r[df_r["filter_status"] != "PASS"]
    fail_filters = ", ".join(df_fail["variant"]) if not df_fail.empty else ""

    return trad_r, trad_s, var_str, comentarios, hetero, borderline, fail_filters, trad

# ============================================================
# GENERATE WORD REPORT
# ============================================================

def gerar_laudo_word(biosample, input_table):

    df = read_input_table(input_table)
    row = df[df["Biosample"] == biosample].iloc[0]

    FARMACOS_R, FARMACOS_S, VAR_RESISTENCIA, COMENTARIOS, hetero, bl, flt, trad = process_resistance(biosample)

    obs = ["Observações:"]
    if hetero:
        obs.append("ʰ Heterorresistência: " + ", ".join(hetero))
    if bl:
        obs.append("† Mutações do tipo borderline: " + ", ".join(bl))
    if flt:
        obs.append("Falharam nos filtros de qualidade: " + flt)

    try:
        df_cov = pd.read_excel(QC_SUMMARY, sheet_name="tbdrRCov", dtype=str)
        row_cov = df_cov[df_cov["biosample"] == biosample]
        if not row_cov.empty:
            cov = row_cov.iloc[0]["Cov < 10"].strip()
            if cov:
                cov_trad = traduzir_farmaco_cov(cov, trad)
                obs.append("Falharam em cobertura do sequenciamento: " + cov_trad)
    except:
        pass

    OBSERVACOES = "\n".join(obs)

    doc = Document(MODEL)

    bold_placeholder(doc, "RESISTENTE")
    bold_placeholder(doc, "FARMACOS_R")
    bold_placeholder(doc, "AMOSTRA")

    mapping = {
	    "REQUISICAO": row["Requisição - Request ID"],
	    "PACIENTE": row["Paciente - Patient Name"],
	    "REQUISITANTE": row["Requisitante - Requesting Clinician"],
	    "ORIGEM": row["Origem - Referring Institution"],
	    "CARTAO_SUS": row["Cartão Nacional de Saúde - National Health ID"],
	    "MUNICIPIO": row["Município - City"],
	    "CADASTRO": fmt_date(row["Data de Cadastro - Registration Date"]),
	    "IDADE": row["Idade - Age"],
	    "SEXO": row["Sexo - Sex"],
	    "PROFISSIONAL": row["Profissional de Saúde - Healthcare Professional"],

	    "REGISTRO": row["Registro Interno - Internal Record ID"],
	    "COLETA": fmt_date(row["Data da Coleta - Collection Date"]),
	    "RECEBIMENTO": fmt_date(row["Data do recebimento - Receipt Date"]),
	    "AMOSTRA": row["Amostra - Sample ID"],
	    "LINHAGEM": get_lineage(biosample),

	    "RESISTENTE": "RESISTENTE",
	    "FARMACOS_R": FARMACOS_R,
	    "VAR_RESISTENCIA": VAR_RESISTENCIA,
	    "COMENTARIOS": COMENTARIOS,

	    "SENSIVEL": "SENSÍVEL",
	    "FARMACOS_S": FARMACOS_S,

	    "HETERORESISTENCIA": OBSERVACOES,

	    # MUST EXIST TO AVOID PLACEHOLDER LEAKAGE
	    "BORDERLINE": "",
	    "FILTROS": "",
	    "COBERTURA": "",

	    "RESPONSAVEL": row["Nome RT - RT Name"],
	    "RGRT": row["Registro RT - RT License Number"],
	    "CONFERENCIA": datetime.now().strftime("%d/%m/%Y"),
	    "EXECUTADO": row["Laboratório responsável - Responsible Laboratory"],
    }


    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, mapping)

    for table in doc.tables:
        replace_in_table(table, mapping)

    replace_header_footer(doc, mapping)

    if FARMACOS_R.strip() == "":
        tabela = doc.tables[2]
        tbl = tabela._tbl
        tbl.remove(tbl.tr_lst[1])

    os.makedirs(OUT_DIR, exist_ok=True)
    out = f"{OUT_DIR}/{biosample}.docx"
    doc.save(out)
    print(f"[OK] Laudo gerado em: {out}")

# ============================================================
# MAIN
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Generate clinical TB resistance Word report."
    )
    parser.add_argument("biosample", help="Biosample ID")
    parser.add_argument(
        "--table",
        default=INPUT_TABLE,
        help="Path to input_table.csv or input_table.xlsx"
    )

    args = parser.parse_args()

    gerar_laudo_word(args.biosample, args.table)

if __name__ == "__main__":
    main()
